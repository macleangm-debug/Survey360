"""DataPulse - Report Builder Routes
Generates professional reports combining narratives, tables, and charts
Supports PDF, Word, HTML export formats
"""

from fastapi import APIRouter, HTTPException, Request, BackgroundTasks
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from datetime import datetime, timezone
import uuid
import io
import json
import os

router = APIRouter(prefix="/reports", tags=["Reports"])


# ============ Models ============

class ReportSection(BaseModel):
    type: str  # "text", "table", "chart", "page_break"
    title: Optional[str] = None
    content: Optional[str] = None  # For text sections (markdown supported)
    data: Optional[Dict[str, Any]] = None  # For tables/charts
    settings: Optional[Dict[str, Any]] = None  # Style settings


class ReportTemplate(BaseModel):
    name: str
    description: Optional[str] = None
    sections: List[ReportSection]
    settings: Optional[Dict[str, Any]] = None  # Global report settings


class CreateReportRequest(BaseModel):
    org_id: str
    title: str
    subtitle: Optional[str] = None
    author: Optional[str] = None
    form_id: Optional[str] = None
    snapshot_id: Optional[str] = None
    sections: List[ReportSection]
    template_id: Optional[str] = None
    settings: Optional[Dict[str, Any]] = None


class GenerateReportRequest(BaseModel):
    report_id: str
    format: str = "pdf"  # pdf, docx, html
    include_appendix: bool = True
    include_methodology: bool = True


# ============ Report Templates ============
# Note: Template routes must be defined before the /{org_id} routes to avoid conflicts

@router.get("/templates/{org_id}")
async def list_templates(request: Request, org_id: str):
    """List available report templates"""
    db = request.app.state.db
    
    templates = await db.report_templates.find(
        {"$or": [{"org_id": org_id}, {"org_id": None}]},  # Org-specific + global
        {"_id": 0}
    ).to_list(50)
    
    # Add default templates if none exist
    if not templates:
        templates = [
            {
                "id": "default-summary",
                "name": "Summary Report",
                "description": "Basic summary with key findings",
                "sections": [
                    {"type": "text", "title": "Executive Summary", "content": ""},
                    {"type": "table", "title": "Key Metrics"},
                    {"type": "text", "title": "Conclusions", "content": ""}
                ]
            },
            {
                "id": "default-detailed",
                "name": "Detailed Analysis Report",
                "description": "Comprehensive report with methodology",
                "sections": [
                    {"type": "text", "title": "Introduction", "content": ""},
                    {"type": "text", "title": "Methodology", "content": ""},
                    {"type": "table", "title": "Descriptive Statistics"},
                    {"type": "chart", "title": "Distribution Charts"},
                    {"type": "page_break"},
                    {"type": "table", "title": "Inferential Statistics"},
                    {"type": "text", "title": "Discussion", "content": ""},
                    {"type": "text", "title": "Conclusions", "content": ""}
                ]
            }
        ]
    
    return templates


@router.post("/templates")
async def create_template(request: Request, template: ReportTemplate, org_id: str):
    """Create a new report template"""
    db = request.app.state.db
    
    template_id = str(uuid.uuid4())
    
    template_doc = {
        "id": template_id,
        "org_id": org_id,
        "name": template.name,
        "description": template.description,
        "sections": [s.dict() for s in template.sections],
        "settings": template.settings or {},
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.report_templates.insert_one(template_doc)
    
    return {"id": template_id, "message": "Template created successfully"}


# ============ Report Management ============

@router.post("")
async def create_report(request: Request, req: CreateReportRequest):
    """Create a new report definition"""
    db = request.app.state.db
    
    report_id = str(uuid.uuid4())
    
    report = {
        "id": report_id,
        "org_id": req.org_id,
        "title": req.title,
        "subtitle": req.subtitle,
        "author": req.author,
        "form_id": req.form_id,
        "snapshot_id": req.snapshot_id,
        "sections": [s.dict() for s in req.sections],
        "template_id": req.template_id,
        "settings": req.settings or {},
        "status": "draft",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.reports.insert_one(report)
    
    return {"id": report_id, "message": "Report created successfully"}


@router.get("/{org_id}")
async def list_reports(request: Request, org_id: str):
    """List all reports for an organization"""
    db = request.app.state.db
    
    reports = await db.reports.find(
        {"org_id": org_id},
        {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    
    return reports


@router.get("/{org_id}/{report_id}")
async def get_report(request: Request, org_id: str, report_id: str):
    """Get a specific report"""
    db = request.app.state.db
    
    report = await db.reports.find_one(
        {"id": report_id, "org_id": org_id},
        {"_id": 0}
    )
    
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return report


@router.put("/{org_id}/{report_id}")
async def update_report(request: Request, org_id: str, report_id: str, req: CreateReportRequest):
    """Update a report"""
    db = request.app.state.db
    
    result = await db.reports.update_one(
        {"id": report_id, "org_id": org_id},
        {"$set": {
            "title": req.title,
            "subtitle": req.subtitle,
            "author": req.author,
            "sections": [s.dict() for s in req.sections],
            "settings": req.settings or {},
            "updated_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    if result.matched_count == 0:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return {"message": "Report updated successfully"}


@router.delete("/{org_id}/{report_id}")
async def delete_report(request: Request, org_id: str, report_id: str):
    """Delete a report"""
    db = request.app.state.db
    
    result = await db.reports.delete_one({"id": report_id, "org_id": org_id})
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return {"message": "Report deleted successfully"}


# ============ Report Generation ============

@router.post("/generate")
async def generate_report(request: Request, req: GenerateReportRequest):
    """Generate a report in the specified format"""
    db = request.app.state.db
    
    # Get report definition
    report = await db.reports.find_one({"id": req.report_id}, {"_id": 0})
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Get associated data if needed
    snapshot_data = None
    form_data = None
    
    if report.get("snapshot_id"):
        snapshot_data = await db.snapshots.find_one(
            {"id": report["snapshot_id"]},
            {"_id": 0}
        )
    
    if report.get("form_id"):
        form_data = await db.forms.find_one(
            {"id": report["form_id"]},
            {"_id": 0}
        )
    
    # Generate report content
    if req.format == "html":
        content = generate_html_report(report, snapshot_data, form_data, req)
        return StreamingResponse(
            io.BytesIO(content.encode()),
            media_type="text/html",
            headers={
                "Content-Disposition": f"attachment; filename=report_{report['id'][:8]}.html"
            }
        )
    
    elif req.format == "docx":
        content = generate_docx_report(report, snapshot_data, form_data, req)
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=report_{report['id'][:8]}.docx"
            }
        )
    
    elif req.format == "pdf":
        html_content = generate_html_report(report, snapshot_data, form_data, req)
        pdf_content = generate_pdf_from_html(html_content)
        return StreamingResponse(
            io.BytesIO(pdf_content),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=report_{report['id'][:8]}.pdf"
            }
        )
    
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {req.format}")


def generate_html_report(report: dict, snapshot: dict, form: dict, req: GenerateReportRequest) -> str:
    """Generate HTML report content"""
    import markdown2
    
    settings = report.get("settings", {})
    theme = settings.get("theme", "light")
    
    # CSS styles
    css = """
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px;
            line-height: 1.6;
            color: #1e293b;
        }
        .report-header {
            text-align: center;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 2px solid #0ea5e9;
        }
        .report-title { font-size: 28px; font-weight: bold; margin-bottom: 10px; }
        .report-subtitle { font-size: 18px; color: #64748b; margin-bottom: 10px; }
        .report-meta { font-size: 14px; color: #94a3b8; }
        .section { margin: 30px 0; }
        .section-title { font-size: 20px; font-weight: 600; margin-bottom: 15px; color: #0369a1; }
        table { width: 100%; border-collapse: collapse; margin: 20px 0; }
        th, td { padding: 12px; text-align: left; border-bottom: 1px solid #e2e8f0; }
        th { background: #f1f5f9; font-weight: 600; }
        tr:hover { background: #f8fafc; }
        .chart-placeholder { 
            background: #f1f5f9; 
            padding: 40px; 
            text-align: center; 
            border-radius: 8px;
            color: #64748b;
        }
        .page-break { page-break-after: always; }
        .methodology {
            background: #f0f9ff;
            padding: 20px;
            border-radius: 8px;
            margin: 20px 0;
        }
        .appendix {
            border-top: 2px solid #e2e8f0;
            margin-top: 40px;
            padding-top: 20px;
        }
        @media print {
            body { padding: 20px; }
            .page-break { page-break-after: always; }
        }
    </style>
    """
    
    # Build HTML content
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{report.get('title', 'Report')}</title>
        {css}
    </head>
    <body>
        <div class="report-header">
            <div class="report-title">{report.get('title', 'Untitled Report')}</div>
            {f'<div class="report-subtitle">{report.get("subtitle")}</div>' if report.get('subtitle') else ''}
            <div class="report-meta">
                {f'Author: {report.get("author")} | ' if report.get('author') else ''}
                Generated: {datetime.now().strftime('%B %d, %Y')}
            </div>
        </div>
    """
    
    # Add sections
    for section in report.get("sections", []):
        if section["type"] == "text":
            content = markdown2.markdown(section.get("content", ""))
            html += f"""
            <div class="section">
                {f'<div class="section-title">{section.get("title")}</div>' if section.get("title") else ''}
                <div class="section-content">{content}</div>
            </div>
            """
        
        elif section["type"] == "table":
            html += render_table_section(section)
        
        elif section["type"] == "chart":
            html += f"""
            <div class="section">
                {f'<div class="section-title">{section.get("title")}</div>' if section.get("title") else ''}
                <div class="chart-placeholder">
                    [Chart: {section.get('data', {}).get('chart_type', 'Unknown')}]
                </div>
            </div>
            """
        
        elif section["type"] == "page_break":
            html += '<div class="page-break"></div>'
    
    # Add methodology section
    if req.include_methodology:
        html += f"""
        <div class="methodology">
            <div class="section-title">Methodology</div>
            <p><strong>Data Source:</strong> {form.get('name', 'Survey Form') if form else 'N/A'}</p>
            <p><strong>Snapshot ID:</strong> {report.get('snapshot_id', 'Live data')}</p>
            <p><strong>Generated:</strong> {datetime.now(timezone.utc).isoformat()}</p>
        </div>
        """
    
    # Add appendix
    if req.include_appendix and snapshot:
        html += f"""
        <div class="appendix">
            <div class="section-title">Appendix: Data Summary</div>
            <p><strong>Total Records:</strong> {len(snapshot.get('data', []))}</p>
            <p><strong>Variables:</strong> {len(snapshot.get('schema', []))}</p>
        </div>
        """
    
    html += """
    </body>
    </html>
    """
    
    return html


def render_table_section(section: dict) -> str:
    """Render a table section as HTML"""
    data = section.get("data", {})
    rows = data.get("rows", [])
    columns = data.get("columns", [])
    
    if not rows or not columns:
        return '<div class="section"><p>No table data available</p></div>'
    
    html = f"""
    <div class="section">
        {f'<div class="section-title">{section.get("title")}</div>' if section.get("title") else ''}
        <table>
            <thead>
                <tr>
                    {''.join(f'<th>{col}</th>' for col in columns)}
                </tr>
            </thead>
            <tbody>
    """
    
    for row in rows[:100]:  # Limit rows
        html += "<tr>"
        for col in columns:
            value = row.get(col, "")
            if isinstance(value, float):
                value = f"{value:.4f}"
            html += f"<td>{value}</td>"
        html += "</tr>"
    
    html += """
            </tbody>
        </table>
    </div>
    """
    
    return html


def generate_docx_report(report: dict, snapshot: dict, form: dict, req: GenerateReportRequest) -> bytes:
    """Generate Word document report"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title
    title = doc.add_heading(report.get("title", "Report"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    if report.get("subtitle"):
        subtitle = doc.add_paragraph(report["subtitle"])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if report.get("author"):
        meta.add_run(f"Author: {report['author']} | ")
    meta.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    
    doc.add_paragraph()  # Spacer
    
    # Sections
    for section in report.get("sections", []):
        if section["type"] == "text":
            if section.get("title"):
                doc.add_heading(section["title"], level=1)
            doc.add_paragraph(section.get("content", ""))
        
        elif section["type"] == "table":
            if section.get("title"):
                doc.add_heading(section["title"], level=1)
            
            data = section.get("data", {})
            rows = data.get("rows", [])
            columns = data.get("columns", [])
            
            if rows and columns:
                table = doc.add_table(rows=1, cols=len(columns))
                table.style = 'Table Grid'
                
                # Header row
                header_cells = table.rows[0].cells
                for i, col in enumerate(columns):
                    header_cells[i].text = str(col)
                
                # Data rows
                for row_data in rows[:50]:
                    row_cells = table.add_row().cells
                    for i, col in enumerate(columns):
                        value = row_data.get(col, "")
                        if isinstance(value, float):
                            value = f"{value:.4f}"
                        row_cells[i].text = str(value)
        
        elif section["type"] == "page_break":
            doc.add_page_break()
    
    # Methodology
    if req.include_methodology:
        doc.add_heading("Methodology", level=1)
        doc.add_paragraph(f"Data Source: {form.get('name', 'Survey Form') if form else 'N/A'}")
        doc.add_paragraph(f"Snapshot ID: {report.get('snapshot_id', 'Live data')}")
        doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).isoformat()}")
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()


def generate_pdf_from_html(html_content: str) -> bytes:
    """Convert HTML-style report to PDF using reportlab"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        import re
        from html.parser import HTMLParser
        
        # Simple HTML parser to extract text content
        class HTMLTextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.sections = []
                self.current_section = None
                self.current_text = ""
                self.in_title = False
                self.in_table = False
                self.table_data = []
                self.current_row = []
                
            def handle_starttag(self, tag, attrs):
                if tag == 'h1':
                    self.in_title = True
                elif tag == 'table':
                    self.in_table = True
                    self.table_data = []
                elif tag == 'tr':
                    self.current_row = []
                    
            def handle_endtag(self, tag):
                if tag == 'h1':
                    self.in_title = False
                    self.sections.append(('title', self.current_text.strip()))
                    self.current_text = ""
                elif tag in ('p', 'div'):
                    if self.current_text.strip():
                        self.sections.append(('text', self.current_text.strip()))
                    self.current_text = ""
                elif tag == 'table':
                    self.in_table = False
                    if self.table_data:
                        self.sections.append(('table', self.table_data))
                    self.table_data = []
                elif tag == 'tr':
                    if self.current_row:
                        self.table_data.append(self.current_row)
                    self.current_row = []
                elif tag in ('td', 'th'):
                    self.current_row.append(self.current_text.strip())
                    self.current_text = ""
                    
            def handle_data(self, data):
                self.current_text += data
        
        # Parse HTML
        parser = HTMLTextExtractor()
        parser.feed(html_content)
        
        # Create PDF
        pdf_buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            pdf_buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#0369a1'),
            spaceAfter=20,
            alignment=TA_CENTER
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#0369a1'),
            spaceAfter=12,
            spaceBefore=20
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            leading=16
        )
        
        # Build content
        story = []
        
        for section_type, content in parser.sections:
            if section_type == 'title':
                story.append(Paragraph(content, title_style))
            elif section_type == 'text':
                # Check if it looks like a heading
                if len(content) < 100 and not content.endswith('.'):
                    story.append(Paragraph(content, heading_style))
                else:
                    story.append(Paragraph(content, body_style))
            elif section_type == 'table':
                if content:
                    t = Table(content)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f1f5f9')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#1e293b')),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#334155')),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 9),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#e2e8f0')),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('TOPPADDING', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                    ]))
                    story.append(Spacer(1, 12))
                    story.append(t)
                    story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        pdf_buffer.seek(0)
        
        return pdf_buffer.getvalue()
    except Exception as e:
        print(f"PDF generation error: {e}")
        # Fallback: return HTML as bytes if PDF generation fails
        return html_content.encode()
